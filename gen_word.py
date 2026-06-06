from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_cn(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    return p

def set_cell_font(cell, text, bold=False, size=Pt(10)):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = size
    run.bold = bold

# ========== 标题 ==========
title = doc.add_heading('任务四  渠道流量推算  计算报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# ========== 一、基本资料 ==========
add_heading_cn('一、基本资料', 1)

add_heading_cn('1.1 灌区基本情况', 2)
items = [
    '总灌溉面积：7.79万亩',
    '土壤类型：中壤土',
    '干、支渠采用续灌工作制度',
    '斗渠、农渠实行轮灌（六支渠为典型渠道：8条斗渠，相邻2个为一轮灌组；12条农渠，相邻4条为一轮灌组）',
]
for item in items:
    add_para('• ' + item, indent=True)

add_heading_cn('1.2 各支渠长度及灌溉面积', 2)
table = doc.add_table(rows=3, cols=8, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['渠道', '一支', '二支', '三支', '四支', '五支', '六支', '总计']
row1 = ['长度(km)', '5.29', '1.8', '2.4', '2.48', '2.4', '2.32', '16.69']
row2 = ['面积(万亩)', '1.39', '1.22', '1.37', '1.4', '1.15', '1.26', '7.79']
for i, h in enumerate(headers):
    set_cell_font(table.rows[0].cells[i], h, bold=True)
    set_cell_font(table.rows[1].cells[i], row1[i])
    set_cell_font(table.rows[2].cells[i], row2[i])

add_heading_cn('1.3 作物种植比例', 2)
crops_ratio = ['棉花：10%', '苜蓿：5%', '林地：12%', '葡萄：11%', '合计：38%（其余为非灌溉面积）']
for item in crops_ratio:
    add_para('• ' + item, indent=True)

# ========== 二、任务1 ==========
add_heading_cn('二、任务1：灌水率图绘制及修正', 1)

add_heading_cn('2.1 灌水率计算公式', 2)
add_formula('q = α × m / (T × 8.64)')
add_para('式中：', indent=True)
add_para('q — 灌水率 [m3/(s·万亩)]', indent=True)
add_para('α — 作物种植比例（小数）', indent=True)
add_para('m — 灌水定额 [m3/亩]', indent=True)
add_para('T — 灌水延续天数 [天]', indent=True)
add_para('8.64 — 换算系数（86400/10000）', indent=True)

add_heading_cn('2.2 各作物各次灌水率计算', 2)

# 棉花
add_para('棉花（α=10%）', bold=True)
cotton_data = [
    ['灌水次序', '灌水时间', '天数T', '定额m(m3/亩)', '灌水率q'],
    ['1', '04-06~04-15', '10', '80', '0.0926'],
    ['2', '06-10~06-20', '11', '60', '0.0631'],
    ['3', '07-01~07-10', '10', '60', '0.0694'],
    ['4', '08-01~08-10', '10', '60', '0.0694'],
    ['5', '09-11~09-20', '10', '60', '0.0694'],
]
t = doc.add_table(rows=len(cotton_data), cols=5, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row in enumerate(cotton_data):
    for j, val in enumerate(row):
        set_cell_font(t.rows[i].cells[j], val, bold=(i==0))

# 苜蓿
add_para('苜蓿（α=5%）', bold=True)
alfalfa_data = [
    ['灌水次序', '灌水时间', '天数T', '定额m(m3/亩)', '灌水率q'],
    ['1', '05-11~05-20', '10', '60', '0.0347'],
    ['2', '06-01~06-10', '10', '60', '0.0347'],
    ['3', '07-11~07-20', '10', '60', '0.0347'],
    ['4', '08-01~08-10', '10', '60', '0.0347'],
    ['5', '08-21~08-31', '11', '60', '0.0316'],
    ['6', '09-11~09-20', '10', '60', '0.0347'],
]
t = doc.add_table(rows=len(alfalfa_data), cols=5, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row in enumerate(alfalfa_data):
    for j, val in enumerate(row):
        set_cell_font(t.rows[i].cells[j], val, bold=(i==0))

# 林地
add_para('林地（α=12%）', bold=True)
forest_data = [
    ['灌水次序', '灌水时间', '天数T', '定额m(m3/亩)', '灌水率q'],
    ['1', '05-11~05-20', '10', '60', '0.0833'],
    ['2', '06-01~06-10', '10', '60', '0.0833'],
    ['3', '07-11~07-20', '10', '60', '0.0833'],
    ['4', '08-01~08-10', '10', '60', '0.0833'],
    ['5', '08-21~08-31', '11', '60', '0.0758'],
    ['6', '09-11~09-20', '10', '60', '0.0833'],
]
t = doc.add_table(rows=len(forest_data), cols=5, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row in enumerate(forest_data):
    for j, val in enumerate(row):
        set_cell_font(t.rows[i].cells[j], val, bold=(i==0))

# 葡萄
add_para('葡萄（α=11%）', bold=True)
grape_data = [
    ['灌水次序', '灌水时间', '天数T', '定额m(m3/亩)', '灌水率q'],
    ['1', '04-16~04-25', '10', '70', '0.0891'],
    ['2', '05-05~05-15', '11', '50', '0.0579'],
    ['3', '06-15~06-25', '11', '50', '0.0579'],
    ['4', '07-08~07-19', '12', '60', '0.0637'],
    ['5', '07-26~08-05', '11', '50', '0.0579'],
    ['6', '08-15~08-25', '11', '50', '0.0579'],
    ['7', '10-18~10-29', '12', '60', '0.0637'],
]
t = doc.add_table(rows=len(grape_data), cols=5, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row in enumerate(grape_data):
    for j, val in enumerate(row):
        set_cell_font(t.rows[i].cells[j], val, bold=(i==0))

add_heading_cn('2.3 修正前灌水率图', 2)
add_para('修正前灌水率图如下所示：', indent=True)
doc.add_picture(r'd:\灌排第四次\灌水率图_修正前.png', width=Inches(6.5))
add_para('修正前主要问题：', indent=True)
add_para('• 08-01~08-05：出现最大峰值 0.2454 m3/(s·万亩)（棉花+林地+苜蓿+葡萄四作物重叠）', indent=True)
add_para('• 07-11~07-19：峰值 0.1817（苜蓿+林地+葡萄重叠）', indent=True)
add_para('• 06-10：峰值 0.1812（棉花+林地+苜蓿重叠）', indent=True)
add_para('• 05-11~05-15：峰值 0.1759（苜蓿+林地+葡萄重叠）', indent=True)

add_heading_cn('2.4 灌水率修正', 2)
add_para('修正原则：', bold=True, indent=True)
add_para('（1）短暂高峰适当削减，将部分灌水时间移至附近低谷', indent=True)
add_para('（2）修正后灌水率应尽量连续平稳', indent=True)
add_para('（3）修正前后总灌水量不变', indent=True)

add_para('修正方案：', bold=True, indent=True)
corr_table = [
    ['修正项', '作物', '原灌水时间', '修正后灌水时间', '说明'],
    ['1', '苜蓿第1次', '05-11~05-20', '05-21~05-30', '避开5月中旬与林地+葡萄重叠高峰'],
    ['2', '葡萄第5次', '07-26~08-05', '07-21~07-31', '避开8月初与棉花+林地+苜蓿重叠高峰'],
    ['3', '苜蓿第4次', '08-01~08-10', '08-11~08-20', '避开8月初多作物重叠高峰'],
]
t = doc.add_table(rows=len(corr_table), cols=5, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row in enumerate(corr_table):
    for j, val in enumerate(row):
        set_cell_font(t.rows[i].cells[j], val, bold=(i==0))

add_para('')
add_para('修正效果：', bold=True, indent=True)
add_para('• 修正前最大灌水率：0.2454 m3/(s·万亩)', indent=True)
add_para('• 修正后最大灌水率：0.1875 m3/(s·万亩)（出现在09-11~09-20）', indent=True)
add_para('• 峰值削减：23.2%', indent=True)

add_heading_cn('2.5 修正后灌水率图', 2)
add_para('修正后灌水率图如下所示：', indent=True)
doc.add_picture(r'd:\灌排第四次\灌水率图_修正后.png', width=Inches(6.5))

add_heading_cn('2.6 设计灌水率确定', 2)
add_para('取修正后灌水率图中的最大值作为设计灌水率：', indent=True)
add_formula('q设计 = 0.1875 m3/(s·万亩)')

# ========== 三、任务2 ==========
add_heading_cn('三、任务2：干渠渠首引水流量计算', 1)

add_heading_cn('3.1 计算公式', 2)
add_formula('Q = q设计 × A / ηf')
add_para('式中：', indent=True)
add_para('Q — 干渠渠首引水流量 [m3/s]', indent=True)
add_para('q设计 — 设计灌水率 [m3/(s·万亩)]', indent=True)
add_para('A — 总灌溉面积 [万亩]', indent=True)
add_para('ηf — 田间水利用系数', indent=True)

add_heading_cn('3.2 计算过程', 2)
add_para('已知：', indent=True)
add_para('• q设计 = 0.1875 m3/(s·万亩)', indent=True)
add_para('• A = 7.79 万亩', indent=True)
add_para('• ηf = 0.90', indent=True)
add_para('')
add_para('代入公式：', indent=True)
add_formula('Q = 0.1875 × 7.79 / 0.90 = 1.4606 / 0.90 = 1.6229 m3/s')

add_heading_cn('3.3 计算结果', 2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Q干渠渠首 = 1.62 m3/s')
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(16)
run.bold = True

# ========== 四、结果汇总 ==========
add_heading_cn('四、结果汇总', 1)
summary = [
    ['项目', '数值'],
    ['设计灌水率 q设计', '0.1875 m3/(s·万亩)'],
    ['总灌溉面积 A', '7.79 万亩'],
    ['田间水利用系数 ηf', '0.90'],
    ['干渠渠首引水流量 Q', '1.62 m3/s'],
]
t = doc.add_table(rows=len(summary), cols=2, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row in enumerate(summary):
    for j, val in enumerate(row):
        set_cell_font(t.rows[i].cells[j], val, bold=(i==0 or (i==4 and j==1)), size=Pt(12))

output_path = r'd:\灌排第四次\任务四_渠道流量推算_计算报告v2.docx'
doc.save(output_path)
print(f'Word文档已保存至: {output_path}')